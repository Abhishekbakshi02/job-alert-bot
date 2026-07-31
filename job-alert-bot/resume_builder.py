"""
Renders resume content (a dict shaped like resume_data.json) into a .docx
file using ALWAYS the same template - only the content differs per job.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_bottom_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "444444")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    _add_bottom_border(p)


def _tab_line(doc, left_text, right_text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.space_after = Pt(2)
    run1 = p.add_run(left_text)
    run1.bold = True
    run1.font.size = Pt(size)
    run2 = p.add_run("\t" + right_text)
    run2.font.size = Pt(size)


def _bullet(doc, text, level=0):
    p = doc.add_paragraph()
    bullet_char = "•  " if level == 0 else "◦  "
    p.paragraph_format.left_indent = Inches(0.25 if level == 0 else 0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(bullet_char + text)
    run.font.size = Pt(9)


def _tech_stack_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    run1 = p.add_run(label + " ")
    run1.bold = True
    run1.font.size = Pt(9)
    run2 = p.add_run(value)
    run2.font.size = Pt(9)


def build_resume_docx(data: dict, output_path: str) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.25)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_p.add_run(data["name"])
    r.bold = True
    r.font.size = Pt(18)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(4)
    contact_p.add_run(
        f"{data['email']}    {data['linkedin']}    {data['phone']}    {data['location']}"
    ).font.size = Pt(9)

    _section_heading(doc, "Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    sp.add_run(data["summary"]).font.size = Pt(9)

    _section_heading(doc, "Technical Skills")
    for skill in data["skills"]:
        _tech_stack_line(doc, skill["label"] + ":", skill["value"])

    _section_heading(doc, "Experience")
    for job in data["experience"]:
        _tab_line(doc, f"{job['title']} | {job['company']}", job["dates"])
        for bullet in job["bullets"]:
            _bullet(doc, bullet, level=0)
        _tech_stack_line(doc, "Tech Stack Used:", job["techStack"])

    _section_heading(doc, "Projects")
    for proj in data["projects"]:
        _tab_line(doc, proj["name"], proj["dates"])
        for bullet in proj["bullets"]:
            _bullet(doc, bullet, level=1)
        _tech_stack_line(doc, "Tech Stack:", proj["techStack"])

    _section_heading(doc, "Education")
    edu = data["education"]
    _tab_line(doc, edu["school"], edu["dates"])
    edu_p = doc.add_paragraph()
    edu_p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    edu_p.add_run(edu["degree"]).font.size = Pt(9)
    edu_p.add_run("\t" + edu["gpa"]).font.size = Pt(9)

    doc.save(output_path)
